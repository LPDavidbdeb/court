import html
import re
from datetime import datetime
from django.views.generic import ListView, DetailView, CreateView, View, FormView
from django.views.generic.edit import UpdateView
from django.urls import reverse_lazy, reverse
from django.shortcuts import redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.conf import settings
from django.contrib import messages
from django.contrib.contenttypes.models import ContentType
from collections import defaultdict
import json
import google.generativeai as genai
import docx
import io
from django.utils.html import strip_tags
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.views.decorators.http import require_POST

from .models import LegalCase, PerjuryContestation, AISuggestion, ExhibitRegistry
from .forms import LegalCaseForm, PerjuryContestationForm, PerjuryContestationNarrativeForm, PerjuryContestationStatementsForm
from .services import refresh_case_exhibits, rebuild_produced_exhibits
from ai_services.utils import EvidenceFormatter
from ai_services.services import analyze_for_json_output, run_police_investigator_service, AI_PERSONAS
from document_manager.models import LibraryNode, DocumentSource, Statement

@require_POST
def update_contestation_title_ajax(request, pk):
    try:
        contestation = get_object_or_404(PerjuryContestation, pk=pk)
        data = json.loads(request.body)
        new_title = data.get('title', '').strip()

        if not new_title:
            return JsonResponse({'status': 'error', 'message': 'Title cannot be empty.'}, status=400)

        contestation.title = new_title
        contestation.save(update_fields=['title'])
        return JsonResponse({'status': 'success', 'new_title': new_title})

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON.'}, status=400)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def retry_parse_suggestion(request, suggestion_pk):
    suggestion = get_object_or_404(AISuggestion, pk=suggestion_pk)
    if not suggestion.raw_response:
        messages.error(request, "No raw response to parse.")
        return redirect('case_manager:contestation_detail', pk=suggestion.contestation.pk)

    try:
        json_match = re.search(r'\{.*\}', suggestion.raw_response, re.DOTALL)
        if json_match:
            cleaned_text = json_match.group(0)
            data_dict = json.loads(cleaned_text)
            
            suggestion.content = data_dict
            suggestion.parsing_success = True
            suggestion.save()
            messages.success(request, "Successfully parsed the raw AI response.")
        else:
            messages.warning(request, "No JSON object could be found in the raw response.")

    except json.JSONDecodeError as e:
        messages.error(request, f"Failed to parse JSON: {e}")
    
    return redirect('case_manager:contestation_detail', pk=suggestion.contestation.pk)

def _get_allegation_context(targeted_statements):
    """
    Helper function to build the enriched text for allegations,
    grouping them by document and including the solemn declaration.
    """
    lies_text = "--- DÉCLARATIONS SOUS SERMENT (VERSION SUSPECTE) ---\n"
    statement_ids = [s.id for s in targeted_statements]

    stmt_content_type = ContentType.objects.get_for_model(Statement)
    nodes = LibraryNode.objects.filter(
        content_type=stmt_content_type,
        object_id__in=statement_ids,
        document__source_type=DocumentSource.REPRODUCED
    ).select_related('document', 'document__author')

    # Group statements by document
    doc_to_stmts = defaultdict(list)
    for node in nodes:
        doc_to_stmts[node.document].append(node.content_object)

    # Format the output
    for doc, stmts in doc_to_stmts.items():
        if doc.solemn_declaration:
            lies_text += f"CONTEXTE DU DOCUMENT : « {doc.title} »\n"
            lies_text += f"DÉCLARATION SOLENNELLE : « {doc.solemn_declaration} »\n\n"
        
        for stmt in stmts:
            author_name = "Auteur Inconnu"
            author_role = ""
            if doc.author:
                author_name = doc.author.get_full_name()
                author_role = f" [{doc.author.role}]"

            doc_date = "Date Inconnue"
            if doc.document_original_date:
                doc_date = doc.document_original_date.strftime('%d %B %Y')
            
            lies_text += f"[ {author_name}{author_role}, dans le document {doc.title} en date du {doc_date} ecrit : « {stmt.text} » ]\n\n"
    
    return lies_text

def serialize_evidence(evidence_pool):
    serialized_data = []
    if evidence_pool.get('events'):
        event_menu = []
        for event in evidence_pool['events']:
            event_menu.append({
                'title': f'Event: {event.date} - {event.explanation[:50]}...',
                'value': f'<blockquote><p>{event.explanation}</p><cite>Event on {event.date}</cite></blockquote>'
            })
        if event_menu:
            serialized_data.append({'title': 'Events', 'menu': event_menu})
    if evidence_pool.get('emails'):
        email_menu = []
        for quote in evidence_pool['emails']:
            email_menu.append({
                'title': f'Email Quote: {quote.quote_text[:50]}...',
                'value': f'<blockquote><p>{quote.quote_text}</p><cite>Email from {quote.email.sender} on {quote.email.date_sent}</cite></blockquote>'
            })
        if email_menu:
            serialized_data.append({'title': 'Email Quotes', 'menu': email_menu})
    return json.dumps(serialized_data)

def preview_ai_context(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    
    refresh_case_exhibits(contestation.case.pk)
    exhibit_registry = contestation.case.exhibits.all()
    exhibit_map = {
        (ex.content_type_id, ex.object_id): ex.get_label()
        for ex in exhibit_registry
    }

    evidence_data = EvidenceFormatter.collect_global_evidence(
        contestation.supporting_narratives.all()
    )

    lies_text = _get_allegation_context(contestation.targeted_statements.all())
    
    prompt_sequence = [
        f"""
        === 1. CADRE DE LA MISSION : RAPPORT DE DÉNONCIATION ===
        RÔLE : Enquêteur spécialisé en fraude judiciaire.
        OBJECTIF : Démontrer l'intention de tromper le tribunal.
        
        {lies_text}

        === 2. THÈSES ET CONTEXTE (Arguments Narratifs) ===
        Voici les dimensions factuelles qui contredisent l'allégation :
        """
    ]

    for i, summary in enumerate(evidence_data['summaries'], 1):
        text_content = strip_tags(summary)
        text_content = html.unescape(text_content)
        text_content = text_content.strip()
        prompt_sequence.append(f"DIMENSION {i} : {text_content}\n")

    prompt_sequence.append("\n=== 3. CHRONOLOGIE UNIFIÉE DES FAITS (PREUVE DIRECTE) ===")
    
    for item in evidence_data['timeline']:
        obj_to_label = item.get('parent_doc', item['obj'])
        label = EvidenceFormatter.get_label(obj_to_label, exhibit_map)
        line = EvidenceFormatter.format_timeline_item(item, exhibit_label=label)
        prompt_sequence.append(line)

    prompt_sequence.append("\n=== 4. INDEX DES PIÈCES & CONTEXTE (DÉTAILS COMPLETS) ===")
    
    def natural_keys(text):
        return [ int(c) if c.isdigit() else c for c in re.split(r'(\d+)', text) ]

    sorted_docs = sorted(
        list(evidence_data['unique_documents']), 
        key=lambda d: natural_keys(EvidenceFormatter.get_label(d, exhibit_map) or "")
    )

    for doc in sorted_docs:
        label = EvidenceFormatter.get_label(doc, exhibit_map)
        context_block = EvidenceFormatter.format_document_reference(doc, exhibit_label=label)
        prompt_sequence.append(context_block + "\n")

    prompt_sequence.append("""
    --- FIN DU DOSSIER ---

    TÂCHE DE RÉDACTION (FORMAT JSON STRICT) :
    Rédige le rapport en 4 sections distinctes.

    CONSIGNES DE TON (CRITIQUE) :
    1. TON : Clinique, factuel, froid et chirurgical.
    2. MISE EN PAGE : Utilise impérativement des sauts de ligne avant chaque point d'une liste numérotée (1., 2., etc.). Aère le texte.
    
    STRUCTURE REQUISE (JSON keys) :
    - suggestion_sec1 (La Déclaration) : Cite la phrase exacte du serment qui est fausse.
    - suggestion_sec2 (La Preuve Contraire) : Résume les faits de la chronologie qui prouvent physiquement l'impossibilité de la déclaration.
    - suggestion_sec3 (Mens Rea / Connaissance) : Démontre que le sujet *savait* que c'était faux au moment de signer (ex: il a lui-même écrit l'email contradictoire P-X).
    - suggestion_sec4 (L'Intention) : Explique *pourquoi* ce mensonge a été fait (le gain judiciaire espéré).
    """)
    
    full_preview = "".join(prompt_sequence)
    return HttpResponse(full_preview, content_type="text/plain; charset=utf-8")

def preview_police_prompt(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    narratives = contestation.supporting_narratives.all()
    
    xml_context = EvidenceFormatter.format_police_context_xml(narratives)
    
    persona = AI_PERSONAS['police_investigator']
    full_prompt = f"{persona['prompt']}\n\n{xml_context}"
    
    return HttpResponse(full_prompt, content_type="text/plain; charset=utf-8")

def generate_exhibit_production(request, pk):
    """
    Trigger to rebuild the ProducedExhibit table.
    """
    case = get_object_or_404(LegalCase, pk=pk)
    
    try:
        count = rebuild_produced_exhibits(case.pk)
        messages.success(request, f"Table des pièces générée avec succès ({count} entrées).")
    except Exception as e:
        messages.error(request, f"Erreur lors de la génération : {str(e)}")
        
    return redirect('case_manager:case_detail', pk=pk)

class LegalCaseListView(ListView):
    model = LegalCase
    template_name = 'case_manager/legalcase_list.html'
    context_object_name = 'cases'
    ordering = ['-created_at']

class LegalCaseDetailView(DetailView):
    model = LegalCase
    template_name = 'case_manager/legalcase_detail.html'
    context_object_name = 'case'
    def get(self, request, *args, **kwargs):
        refresh_case_exhibits(self.kwargs['pk'])
        return super().get(request, *args, **kwargs)
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['contestations'] = self.object.contestations.all()
        context['produced_exhibits'] = self.object.produced_exhibits.all()
        return context

class LegalCaseCreateView(CreateView):
    model = LegalCase
    form_class = LegalCaseForm
    template_name = 'case_manager/legalcase_form.html'
    def get_success_url(self):
        return reverse_lazy('case_manager:case_detail', kwargs={'pk': self.object.pk})

class LegalCaseExportView(View):
    def get(self, request, *args, **kwargs):
        case = get_object_or_404(LegalCase, pk=self.kwargs['pk'])
        document = docx.Document()

        # ------------------------------------------------------------------
        # HELPER: Get precise datetime for sorting exhibits
        # ------------------------------------------------------------------
        def get_datetime_for_sorting(exhibit):
            obj = exhibit.content_object
            dt = None
            if hasattr(obj, 'date_sent') and obj.date_sent:
                dt = obj.date_sent
            elif hasattr(obj, 'date') and obj.date:
                dt = datetime.combine(obj.date, datetime.min.time())
            elif hasattr(obj, 'document_original_date') and obj.document_original_date:
                dt = datetime.combine(obj.document_original_date, datetime.min.time())
            elif hasattr(obj, 'created_at') and obj.created_at:
                dt = obj.created_at
            
            if dt and timezone.is_naive(dt):
                return timezone.make_aware(dt, timezone.get_current_timezone())
            
            return dt or timezone.now()

        # ------------------------------------------------------------------
        # HELPER: Add internal hyperlink in docx
        # ------------------------------------------------------------------
        def add_hyperlink(paragraph, text, anchor):
            part = paragraph.part
            r_id = part.relate_to(anchor, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
            hyperlink.set(docx.oxml.shared.qn('w:anchor'), anchor, )
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            r = paragraph.add_run()
            r._r.append(hyperlink)
            r.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
            r.font.underline = True
            return hyperlink

        # ------------------------------------------------------------------
        # STEP 1: ON-THE-FLY RENUMBERING
        # ------------------------------------------------------------------
        all_exhibits = case.exhibits.all().select_related('content_type')
        sorted_exhibits = sorted(all_exhibits, key=get_datetime_for_sorting)

        renumbering_map = {}
        bookmark_map = {}
        for i, exhibit in enumerate(sorted_exhibits, 1):
            old_label = exhibit.get_label()
            new_label = f"P-{i}"
            renumbering_map[old_label] = new_label
            bookmark_map[old_label] = f"exhibit_{i}"

        def update_references(text, mapping):
            if not text: return ""
            # This regex ensures we only match P- followed by digits
            pattern = re.compile(r'\b(P-\d+)\b')
            
            def replacer(match):
                old_ref = match.group(1)
                return mapping.get(old_ref, old_ref) # Replace or keep original if not in map
            
            return pattern.sub(replacer, text)

        # ------------------------------------------------------------------
        # HELPER: Text Cleaning & Markdown Parsing
        # ------------------------------------------------------------------
        def clean_text(text):
            if not text: return ""
            text = text.replace('</p>', '\n').replace('<br>', '\n').replace('<br/>', '\n')
            text = strip_tags(text)
            text = html.unescape(text)
            return text.strip()

        def add_markdown_content(doc, raw_text, renumbering_map):
            # Update references before cleaning and adding to doc
            updated_text = update_references(raw_text, renumbering_map)
            text = clean_text(updated_text)
            if not text: return

            text = re.sub(r'([\.\:\;])\s+([\*\-]\s)', r'\1\n\2', text)
            text = re.sub(r'([\.\:\;])\s+(\d+\.\s)', r'\1\n\2', text)
            text = re.sub(r'(»)\s+([\*\-\d])', r'\1\n\2', text)
            text = text.replace('***', '* **')

            lines = text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line: continue

                para_style = None
                if re.match(r'^[\*\-]\s+', line):
                    para_style = 'List Bullet'
                    line = re.sub(r'^[\*\-]\s+', '', line)
                elif re.match(r'^\d+\.\s+', line):
                    para_style = 'List Number'
                    line = re.sub(r'^\d+\.\s+', '', line)

                p = doc.add_paragraph(style=para_style)
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        clean_part = part[2:-2]
                        if clean_part:
                            p.add_run(clean_part).bold = True
                    else:
                        p.add_run(part)

        # ------------------------------------------------------------------
        # DOCUMENT GENERATION (Margins & Structure)
        # ------------------------------------------------------------------
        section = document.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        document.add_heading(f'Dénonciation: {case.title}', level=0)
        
        for contestation in case.contestations.all():
            document.add_heading(contestation.title, level=2)
            
            document.add_heading('1. Déclaration', level=3)
            add_markdown_content(document, contestation.final_sec1_declaration, renumbering_map)
            
            document.add_heading('2. Preuve', level=3)
            add_markdown_content(document, contestation.final_sec2_proof, renumbering_map)
            
            document.add_heading('3. Mens Rea', level=3)
            add_markdown_content(document, contestation.final_sec3_mens_rea, renumbering_map)
            
            document.add_heading('4. Intention', level=3)
            add_markdown_content(document, contestation.final_sec4_intent, renumbering_map)
            
            document.add_page_break()

        # --- TABLE OF EXHIBITS (CHRONOLOGICAL) ---
        document.add_heading('Index des Pièces (Exhibits)', level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cote'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Date'

        for exhibit in sorted_exhibits:
            row_cells = table.add_row().cells
            
            # Cell 0: Add hyperlink to bookmark
            old_label = exhibit.get_label()
            new_label = renumbering_map[old_label]
            bookmark_name = bookmark_map[old_label]
            add_hyperlink(row_cells[0].paragraphs[0], new_label, bookmark_name)

            # Cell 1: Description
            obj = exhibit.content_object
            row_cells[1].text = str(obj)

            # Cell 2: Date
            exhibit_date = get_datetime_for_sorting(exhibit)
            row_cells[2].text = exhibit_date.strftime('%Y-%m-%d %H:%M') if isinstance(exhibit_date, datetime) else exhibit_date.strftime('%Y-%m-%d')

        # ==================================================================
        # ANNEXES (CHRONOLOGICAL)
        # ==================================================================
        document.add_page_break()
        document.add_heading('ANNEXES - CONTENU DÉTAILLÉ', level=0)

        for exhibit in sorted_exhibits:
            obj = exhibit.content_object
            old_label = exhibit.get_label()
            new_label = renumbering_map[old_label]
            bookmark_name = bookmark_map[old_label]
            model_name = exhibit.content_type.model

            # Add heading with bookmark
            heading_paragraph = document.add_heading(f'Pièce {new_label}', level=1)
            # This is a simplified way to add a bookmark
            bookmark_start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
            bookmark_start.set(docx.oxml.shared.qn('w:id'), '0')
            bookmark_start.set(docx.oxml.shared.qn('w:name'), bookmark_name)
            heading_paragraph._p.insert(0, bookmark_start)
            bookmark_end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
            bookmark_end.set(docx.oxml.shared.qn('w:id'), '0')
            heading_paragraph._p.append(bookmark_end)

            # --- Content generation for each exhibit type ---
            if model_name == 'email':
                p = document.add_paragraph()
                p.add_run(f"Date : {obj.date_sent}\n").bold = True
                p.add_run(f"De : {obj.sender}\n").bold = True
                p.add_run(f"À : {obj.recipients_to}\n").bold = True
                p.add_run(f"Sujet : {obj.subject}").bold = True
                document.add_paragraph('--- Contenu ---').italic = True
                
                raw_body = obj.body_plain_text or "[Vide]"
                body_lines = raw_body.splitlines()
                cleaned_lines = [line for line in body_lines if not line.strip().startswith('>')]
                cleaned_body = "\n".join(cleaned_lines)
                body_text = clean_text(cleaned_body)
                document.add_paragraph(body_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            elif model_name == 'event':
                document.add_paragraph(f"Date : {obj.date}")
                p = document.add_paragraph()
                p.add_run("Description : ").bold = True
                add_markdown_content(document, obj.explanation, renumbering_map)

                photos = obj.linked_photos.all()
                if photos.exists():
                    document.add_paragraph("Preuve visuelle :").italic = True
                    photo_table = document.add_table(rows=0, cols=2)
                    row_cells = None
                    for index, photo in enumerate(photos):
                        if index % 2 == 0:
                            row_cells = photo_table.add_row().cells
                        cell = row_cells[index % 2]
                        if photo.file:
                            try:
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(photo.file.open(), width=Inches(2.8))
                                caption = cell.add_paragraph(photo.file_name or "Image")
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                cell.add_paragraph(f"[Erreur: {e}]")

            elif model_name == 'photodocument':
                document.add_paragraph(f"Titre : {obj.title}")
                if obj.description:
                    add_markdown_content(document, obj.description, renumbering_map)
                if hasattr(obj, 'ai_analysis') and obj.ai_analysis:
                    document.add_heading("Analyse IA :", level=4)
                    add_markdown_content(document, obj.ai_analysis, renumbering_map)
                
                photos = obj.photos.all()
                if photos.exists():
                    photo_table = document.add_table(rows=0, cols=2)
                    row_cells = None
                    for index, photo in enumerate(photos):
                        if index % 2 == 0:
                            row_cells = photo_table.add_row().cells
                        cell = row_cells[index % 2]
                        if photo.file:
                            try:
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(photo.file.open(), width=Inches(2.8))
                                caption = cell.add_paragraph(f"Page {index + 1}")
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                cell.add_paragraph(f"[Erreur: {e}]")

            elif model_name == 'pdfdocument':
                document.add_paragraph(f"Document : {obj.title}")
                if hasattr(obj, 'ai_analysis') and obj.ai_analysis:
                    document.add_heading("Résumé / Analyse :", level=3)
                    add_markdown_content(document, obj.ai_analysis, renumbering_map)
                document.add_paragraph("[Voir fichier PDF joint]").italic = True

            document.add_page_break()

        # Save
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.getvalue(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="case_{case.pk}_export.docx"'
        return response

class PoliceComplaintExportView(View):
    def get(self, request, *args, **kwargs):
        case = get_object_or_404(LegalCase, pk=self.kwargs['pk'])
        document = docx.Document()
        
        style = document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        document.add_heading(f"DOSSIER DE PLAINTE : {case.title}", level=0)
        document.add_paragraph(f"Date du rapport : {timezone.now().strftime('%Y-%m-%d')}")
        document.add_paragraph("À l'attention des enquêteurs.")
        document.add_page_break()

        contestations = case.contestations.exclude(police_report_data={})

        if not contestations.exists():
            document.add_paragraph("Aucune plainte policière n'a été générée pour ce dossier.")

        for contestation in contestations:
            data = contestation.police_report_data
            
            titre = data.get('titre_document', f"PLAINTE - {contestation.title}")
            document.add_heading(titre, level=1)
            
            sections = data.get('sections', [])
            for section in sections:
                if 'titre' in section:
                    document.add_heading(section['titre'], level=2)
                
                if 'contenu' in section:
                    content = section['contenu']
                    if isinstance(content, list):
                        for item in content:
                            p = document.add_paragraph(str(item))
                            p.style = 'List Bullet'
                    else:
                        document.add_paragraph(str(content))
            
            document.add_page_break()

        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="PLAINTE_POLICE_{case.pk}.docx"'
        return response

class PerjuryContestationCreateView(CreateView):
    model = PerjuryContestation
    form_class = PerjuryContestationForm
    template_name = 'case_manager/perjurycontestation_form.html'
    def form_valid(self, form):
        form.instance.case = get_object_or_404(LegalCase, pk=self.kwargs['case_pk'])
        return super().form_valid(form)
    def get_success_url(self):
        return reverse_lazy('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

class ManageContestationNarrativesView(UpdateView):
    model = PerjuryContestation
    form_class = PerjuryContestationNarrativeForm
    template_name = 'case_manager/manage_narratives.html'
    context_object_name = 'contestation'

    def get_success_url(self):
        messages.success(self.request, "Supporting narratives updated successfully.")
        return reverse('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

class ManageContestationStatementsView(UpdateView):
    model = PerjuryContestation
    form_class = PerjuryContestationStatementsForm
    template_name = 'case_manager/manage_statements.html'
    context_object_name = 'contestation'

    def get_success_url(self):
        messages.success(self.request, "Targeted statements updated successfully.")
        return reverse('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

class PerjuryContestationDetailView(UpdateView):
    model = PerjuryContestation
    template_name = 'case_manager/perjurycontestation_detail.html'
    context_object_name = 'contestation'
    fields = ['final_sec1_declaration', 'final_sec2_proof', 'final_sec3_mens_rea', 'final_sec4_intent']
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        evidence_pool = {'events': [], 'emails': [], 'pdfs': []}
        for narrative in self.object.supporting_narratives.all():
            evidence_pool['events'].extend(narrative.evenements.all())
            evidence_pool['emails'].extend(narrative.citations_courriel.all())
            evidence_pool['pdfs'].extend(narrative.citations_pdf.all())
        context['evidence_json'] = serialize_evidence(evidence_pool)
        context['ai_drafts'] = self.object.ai_suggestions.order_by('-created_at')
        return context

    def get_success_url(self):
        return reverse('case_manager:contestation_detail', kwargs={'pk': self.object.pk})

def generate_ai_suggestion(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    
    narratives_context = []
    
    for narrative in contestation.supporting_narratives.all():
        analysis = narrative.get_structured_analysis()
        
        narrative_block = f"--- TRAME FACTUELLE : {narrative.titre} ---\n"
        
        if 'constats_objectifs' in analysis:
            for constat in analysis['constats_objectifs']:
                narrative_block += f"FAIT ÉTABLI : {constat.get('fait_identifie', 'N/A')}\n"
                narrative_block += f"DÉTAIL : {constat.get('description_factuelle', '')}\n"
                narrative_block += f"IMPACT : {constat.get('contradiction_directe', '')}\n\n"
        else:
            narrative_block += f"RÉSUMÉ MANUEL : {narrative.resume}\n"
            
        narratives_context.append(narrative_block)

    full_evidence_text = "\n".join(narratives_context)

    prompt_sequence = [
        """
        RÔLE : Stratège Juridique Senior (Procureur).
        MISSION : Rédiger un argumentaire de parjure dévastateur basé sur des FAITS VÉRIFIÉS.
        
        TU NE DOIS PAS : Chercher des preuves (c'est déjà fait).
        TU DOIS : Prouver l'INTENTION de mentir (Mens Rea) en connectant les faits.
        """,
        
        f"=== CIBLE (DÉCLARATION SOUS SERMENT) ===\n{_get_allegation_context(contestation.targeted_statements.all())}",
        
        f"=== AUDIT DES FAITS (PREUVE IRRÉFUTABLE) ===\n{full_evidence_text}",
        
        """
        === DIRECTIVES DE RÉDACTION ===
        Rédige le rapport au format JSON strict.
        
        Section 3 (Mens Rea) est la plus importante : Explique comment la multiplicité des faits (les dates, les photos, les emails) prouve qu'il est IMPOSSIBLE que le sujet ait fait une simple "erreur". C'est un mensonge calculé.
        
        Structure JSON attendue :
        {
            "suggestion_sec1": "Citation exacte et contexte...",
            "suggestion_sec2": "Synthèse des faits contraires (utilise les faits de l'audit)...",
            "suggestion_sec3": "Argumentaire sur la Connaissance (Mens Rea)...",
            "suggestion_sec4": "Argumentaire sur l'Intention (Gain judiciaire)..."
        }
        """
    ]
    
    try:
        raw_text = analyze_for_json_output(prompt_sequence)
        
        suggestion = AISuggestion.objects.create(
            contestation=contestation,
            raw_response=raw_text,
            content={},
            parsing_success=False
        )

        try:
            data_dict = json.loads(raw_text)
            suggestion.content = data_dict
            suggestion.parsing_success = True
            suggestion.save()
            messages.success(request, "Stratégie de parjure générée avec succès sur la base de l'audit.")
        except json.JSONDecodeError:
            messages.warning(request, "Réponse générée mais format JSON invalide.")

    except Exception as e:
        messages.error(request, f"Erreur API : {e}")

    return redirect('case_manager:contestation_detail', pk=contestation.pk)

def generate_police_report(request, contestation_pk):
    contestation = get_object_or_404(PerjuryContestation, pk=contestation_pk)
    
    narratives = contestation.supporting_narratives.prefetch_related(
        'evenements', 'citations_courriel', 'citations_pdf', 'photo_documents', 'targeted_statements'
    ).all()
    
    try:
        raw_json = run_police_investigator_service(narratives)
        data = json.loads(raw_json)
        
        contestation.police_report_data = data
        contestation.police_report_date = timezone.now()
        contestation.save()
        
        messages.success(request, "Rapport de police généré avec succès.")
    except Exception as e:
        messages.error(request, f"Erreur lors de la génération : {e}")
        
    return redirect('case_manager:contestation_detail', pk=contestation.pk)
