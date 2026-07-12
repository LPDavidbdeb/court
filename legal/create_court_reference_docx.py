from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parent / "reference_cour_superieure_quebec.docx"


def set_font(style, name: str = "Arial", size: int = 12, bold: bool = False) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = False
    font.color.rgb = RGBColor(0, 0, 0)

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def paragraph_format(
    style,
    *,
    line_spacing: float = 1.5,
    space_before: int = 0,
    space_after: int = 0,
    left_indent_cm: float | None = None,
    right_indent_cm: float | None = None,
    first_line_indent_cm: float | None = None,
    alignment=None,
) -> None:
    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)
    if right_indent_cm is not None:
        fmt.right_indent = Cm(right_indent_cm)
    if first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(first_line_indent_cm)
    if alignment is not None:
        fmt.alignment = alignment


def add_or_get_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def set_table_defaults(document: Document) -> None:
    table_style = document.styles["Table Grid"]
    table_style.font.name = "Arial"
    table_style.font.size = Pt(10)
    table_style.font.color.rgb = RGBColor(0, 0, 0)

    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    table.cell(0, 0).text = "Exemple"
    table.cell(0, 1).text = "Style de tableau"


def build_reference_docx() -> Path:
    document = Document()

    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.left_margin = Cm(3.0)

    normal = document.styles["Normal"]
    set_font(normal, "Arial", 12)
    paragraph_format(normal, line_spacing=1.5, space_before=0, space_after=6)

    title = document.styles["Title"]
    set_font(title, "Arial", 14, bold=True)
    paragraph_format(title, line_spacing=1.5, space_before=0, space_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    title.font.all_caps = True

    for level in range(1, 7):
        heading = document.styles[f"Heading {level}"]
        size = 12 if level > 1 else 14
        set_font(heading, "Arial", size, bold=True)
        paragraph_format(
            heading,
            line_spacing=1.5,
            space_before=12 if level <= 2 else 6,
            space_after=6,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        if level <= 2:
            heading.font.all_caps = True

    body_text = add_or_get_style(document, "Body Text")
    body_text.base_style = normal
    set_font(body_text, "Arial", 12)
    paragraph_format(body_text, line_spacing=1.5, space_before=0, space_after=6)

    first_paragraph = add_or_get_style(document, "First Paragraph")
    first_paragraph.base_style = normal
    set_font(first_paragraph, "Arial", 12)
    paragraph_format(first_paragraph, line_spacing=1.5, space_before=0, space_after=6)

    block_text = add_or_get_style(document, "Block Text")
    block_text.base_style = normal
    set_font(block_text, "Arial", 12)
    paragraph_format(
        block_text,
        line_spacing=1.0,
        space_before=6,
        space_after=6,
        left_indent_cm=1.25,
        right_indent_cm=1.25,
    )

    quote = add_or_get_style(document, "Quote")
    quote.base_style = block_text
    set_font(quote, "Arial", 12)
    paragraph_format(
        quote,
        line_spacing=1.0,
        space_before=6,
        space_after=6,
        left_indent_cm=1.25,
        right_indent_cm=1.25,
    )

    footnote_text = add_or_get_style(document, "Footnote Text")
    set_font(footnote_text, "Arial", 10)
    paragraph_format(footnote_text, line_spacing=1.0, space_before=0, space_after=0)

    for list_style_name in ("List Paragraph", "List Bullet", "List Number"):
        try:
            list_style = document.styles[list_style_name]
        except KeyError:
            continue
        set_font(list_style, "Arial", 12)
        paragraph_format(list_style, line_spacing=1.5, space_before=0, space_after=6, left_indent_cm=0.75)

    for code_style_name in ("Source Code", "Verbatim Char"):
        code_style = add_or_get_style(
            document,
            code_style_name,
            WD_STYLE_TYPE.CHARACTER if code_style_name.endswith("Char") else WD_STYLE_TYPE.PARAGRAPH,
        )
        set_font(code_style, "Courier New", 10)
        if code_style.type == WD_STYLE_TYPE.PARAGRAPH:
            paragraph_format(code_style, line_spacing=1.0, space_before=6, space_after=6, left_indent_cm=1.25)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_paragraph)

    document.core_properties.title = "Modele acte de procedure - Cour superieure du Quebec"
    document.core_properties.subject = "Reference DOCX Pandoc"
    document.core_properties.author = "Court project"

    document.add_heading("DEMANDE INTRODUCTIVE D'INSTANCE", level=1)
    document.add_paragraph(
        "Ce document est un modele de styles pour Pandoc. Son contenu est ignore lors de la conversion.",
        style="Normal",
    )
    document.add_heading("Titre de niveau 2", level=2)
    document.add_paragraph("Paragraphe normal: Arial 12, interligne 1,5, encre noire.", style="Normal")
    document.add_paragraph(
        "Citation ou extrait: Arial 12, interligne simple, retrait gauche et droit de 1,25 cm.",
        style="Block Text",
    )
    set_table_defaults(document)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_reference_docx())
